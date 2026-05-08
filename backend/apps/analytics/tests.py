from io import BytesIO
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.core.management import call_command
from rest_framework import status
from rest_framework.test import APITestCase
from openpyxl import Workbook
from reportlab.pdfgen import canvas

try:
	from docx import Document
except ImportError:
	Document = None

try:
	from pypdf import PdfReader
except ImportError:
	PdfReader = None

from .models import (
	EmisionLote,
	Empresa,
	EspecieMadera,
	FactorEmision,
	HistorialCambioLote,
	Lote,
	UnidadOperativa,
)
from .services.activity_semantics import is_diesel_activity, is_electricity_activity
from .services.decision_engine import calculate_risk_profile, optimize_rows, simulate_rows
from .services.factor_classifier import infer_categoria, normalize_key


class AnalyticsApiIntegrationTest(APITestCase):
	def test_dashboard_endpoint_returns_summary(self):
		response = self.client.get("/api/dashboard/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertIn("total_emisiones", response.data)
		self.assertIn("datos", response.data)
		self.assertIn("emisiones_por_empresa", response.data)

	def test_dashboard_empty_returns_zero_internal_metrics(self):
		response = self.client.get("/api/dashboard/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["source"], "internal")
		self.assertEqual(response.data["total_emisiones"], 0)
		self.assertEqual(response.data["cantidad_registros"], 0)
		self.assertEqual(response.data["empresa_critica"], "Sin datos")
		self.assertEqual(response.data["actividad_critica"], "Sin datos")

	def test_dashboard_with_lote_activities_returns_internal_metrics(self):
		lote_sur = Lote.objects.create(
			id_lote="LOTE-DASH-001",
			empresa_aserradero="Aserradero Sur",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote_norte = Lote.objects.create(
			id_lote="LOTE-DASH-002",
			empresa_aserradero="Aserradero Norte",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=8,
			origen="Temuco",
			destino="Concepcion",
		)
		lote_sur.actividades.create(
			actividad="diesel",
			cantidad=100,
			unidad="litros",
			factor_emision=2.68,
		)
		lote_sur.actividades.create(
			actividad="electricidad",
			cantidad=200,
			unidad="kWh",
			factor_emision=0.41,
		)
		lote_norte.actividades.create(
			actividad="diesel",
			cantidad=50,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/dashboard/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["source"], "internal")
		self.assertAlmostEqual(response.data["total_emisiones"], 484.0)
		self.assertEqual(response.data["cantidad_registros"], 3)
		self.assertEqual(response.data["emisiones_por_empresa"]["Aserradero Sur"], 350.0)
		self.assertEqual(response.data["emisiones_por_empresa"]["Aserradero Norte"], 134.0)
		self.assertEqual(response.data["emisiones_por_actividad"]["diesel"], 402.0)
		self.assertEqual(response.data["emisiones_por_actividad"]["electricidad"], 82.0)
		self.assertEqual(response.data["empresa_critica"], "Aserradero Sur")
		self.assertEqual(response.data["actividad_critica"], "diesel")
		self.assertTrue(response.data["diesel_presente"])
		self.assertGreater(response.data["score_riesgo"], 0)
		self.assertNotEqual(response.data["empresa_critica"], "Sin datos")
		self.assertNotEqual(response.data["actividad_critica"], "Sin datos")

	def test_empresa_unidad_lote_and_activity_hierarchy(self):
		empresa = Empresa.objects.create(empresa_id="EMP-001", nombre="Maderas del Sur")
		unidad = UnidadOperativa.objects.create(
			unidad_id="UNI-SEC-001",
			empresa=empresa,
			nombre="Secado",
			tipo=UnidadOperativa.Tipo.SECADO,
		)
		lote = Lote.objects.create(
			id_lote="LOTE-JER-001",
			empresa=empresa,
			unidad_operativa=unidad,
			empresa_aserradero=empresa.nombre,
			fecha="2026-04-29",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Temuco",
			destino="Coronel",
		)
		actividad_lote = EmisionLote.objects.create(
			lote=lote,
			actividad="Diésel - combustión móvil",
			cantidad=1,
			unidad="m3",
			factor_emision=2740,
		)
		actividad_unidad = EmisionLote.objects.create(
			empresa=empresa,
			unidad_operativa=unidad,
			actividad="Electricidad Los Lagos 2023",
			cantidad=100,
			unidad="kWh",
			factor_emision=0.5,
		)
		actividad_empresa = EmisionLote.objects.create(
			empresa=empresa,
			actividad="Papel virgen",
			cantidad=2,
			unidad="t",
			factor_emision=10,
		)

		self.assertEqual(empresa.unidades_operativas.count(), 1)
		self.assertEqual(unidad.empresa, empresa)
		self.assertEqual(lote.empresa, empresa)
		self.assertEqual(lote.unidad_operativa, unidad)
		self.assertEqual(actividad_lote.tipo_asignacion, EmisionLote.TipoAsignacion.LOTE)
		self.assertEqual(actividad_lote.empresa, empresa)
		self.assertEqual(actividad_lote.unidad_operativa, unidad)
		self.assertEqual(actividad_unidad.tipo_asignacion, EmisionLote.TipoAsignacion.UNIDAD)
		self.assertIsNone(actividad_unidad.lote)
		self.assertEqual(actividad_empresa.tipo_asignacion, EmisionLote.TipoAsignacion.EMPRESA)
		self.assertIsNone(actividad_empresa.lote)
		self.assertIsNone(actividad_empresa.unidad_operativa)

	def test_crear_empresa_crea_unidad_general(self):
		response = self.client.post(
			"/api/empresas/",
			data={
				"empresa_id": "EMP-WORKSPACE",
				"nombre": "Empresa Workspace",
				"rut": "76.123.456-7",
			},
			format="json",
		)

		self.assertEqual(response.status_code, status.HTTP_201_CREATED)
		empresa = Empresa.objects.get(empresa_id="EMP-WORKSPACE")
		self.assertEqual(empresa.unidades_operativas.count(), 1)
		unidad = empresa.unidades_operativas.get()
		self.assertEqual(unidad.nombre, "Unidad General")
		self.assertEqual(unidad.tipo, UnidadOperativa.Tipo.GENERAL)
		self.assertEqual(response.data["unidad_inicial"]["unidad_id"], unidad.unidad_id)

	def test_import_empresas_requires_all_company_columns(self):
		uploaded_file = SimpleUploadedFile(
			"empresas.csv",
			(
				"empresa_id,nombre,rut,direccion\n"
				"EMP-001,Empresa Demo,76.123.456-7,Calle 1\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/empresas/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
		self.assertIn("Faltan columnas obligatorias", response.data["error"])

	def test_import_empresa_completa_rejects_empty_empresa_sheet(self):
		workbook = Workbook()
		for sheet_name, headers in {
			"empresa": [
				"empresa_id",
				"nombre",
				"rut",
				"region",
				"comuna",
				"direccion",
				"rubro",
				"email",
				"telefono",
				"contacto",
				"observaciones",
			],
			"unidades": ["unidad_id", "empresa_id", "nombre", "tipo"],
			"lotes": ["id_lote", "empresa", "fecha", "especie", "volumen_m3", "origen"],
			"actividades": ["id_lote", "actividad", "cantidad", "unidad", "fecha"],
			"factores": ["actividad", "unidad", "factor_emision", "fuente", "anio"],
		}.items():
			sheet = workbook.active if workbook.sheetnames == ["Sheet"] and sheet_name == "empresa" else workbook.create_sheet(title=sheet_name)
			sheet.title = sheet_name
			sheet.append(headers)

		buffer = BytesIO()
		workbook.save(buffer)
		buffer.seek(0)
		uploaded_file = SimpleUploadedFile(
			"empresa_completa_vacia.xlsx",
			buffer.read(),
			content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
		)

		preview_response = self.client.post(
			"/api/importaciones/empresa-completa/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/empresa-completa/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertTrue(preview_response.data["blocking_errors"])
		self.assertTrue(
			any(
				message in error
				for error in preview_response.data["blocking_errors"]
				for message in ["No se encontro una empresa valida", "La hoja empresa no contiene filas validas para importar"]
			)
		)
		self.assertEqual(confirm_response.status_code, status.HTTP_400_BAD_REQUEST)
		self.assertIn("No se pudo confirmar la importación de empresa completa", confirm_response.data["error"])

	def test_empresa_dashboard_aisla_datos_por_empresa(self):
		empresa_a = Empresa.objects.create(empresa_id="EMP-DASH-A", nombre="Empresa A")
		empresa_b = Empresa.objects.create(empresa_id="EMP-DASH-B", nombre="Empresa B")
		unidad_a = UnidadOperativa.objects.create(
			unidad_id="UNI-A",
			empresa=empresa_a,
			nombre="Aserradero A",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		unidad_b = UnidadOperativa.objects.create(
			unidad_id="UNI-B",
			empresa=empresa_b,
			nombre="Aserradero B",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		lote_a = Lote.objects.create(
			id_lote="LOTE-DASH-A",
			empresa=empresa_a,
			unidad_operativa=unidad_a,
			empresa_aserradero=empresa_a.nombre,
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		Lote.objects.create(
			id_lote="LOTE-DASH-B",
			empresa=empresa_b,
			unidad_operativa=unidad_b,
			empresa_aserradero=empresa_b.nombre,
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote_a.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)
		Lote.objects.get(id_lote="LOTE-DASH-B").actividades.create(
			actividad="diesel",
			cantidad=10,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/empresas/EMP-DASH-A/dashboard/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["empresa_id"], "EMP-DASH-A")
		self.assertAlmostEqual(float(response.data["emisiones_totales"]), 214.4)
		self.assertEqual(response.data["lotes_count"], 1)
		self.assertEqual(response.data["actividades_count"], 1)

	def test_empresa_import_lotes_preview_and_confirm_without_empresa_column(self):
		empresa = Empresa.objects.create(empresa_id="EMP-SCOPE-LOTES", nombre="Empresa Scope")
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino\n"
				"LOTE-SCOPE,Empresa Scope,2026-04-28,Pino radiata,10,Curico,Santiago\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			f"/api/empresas/{empresa.empresa_id}/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			f"/api/empresas/{empresa.empresa_id}/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		lote = Lote.objects.get(id_lote="LOTE-SCOPE")
		self.assertEqual(lote.empresa, empresa)

	def test_empresa_import_actividades_rejects_mismatched_company(self):
		empresa_a = Empresa.objects.create(empresa_id="EMP-SCOPE-A", nombre="Scope A")
		empresa_b = Empresa.objects.create(empresa_id="EMP-SCOPE-B", nombre="Scope B")
		unidad_a = UnidadOperativa.objects.create(
			unidad_id="UNI-SCOPE-A",
			empresa=empresa_a,
			nombre="Aserradero Scope A",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		Lote.objects.create(
			id_lote="LOTE-SCOPE-A",
			empresa=empresa_a,
			unidad_operativa=unidad_a,
			empresa_aserradero=empresa_a.nombre,
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			(
				"id_lote,actividad,cantidad,unidad,fecha\n"
				"LOTE-SCOPE-A,diesel,10,litros,2026-04-28\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			f"/api/empresas/{empresa_b.empresa_id}/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["filas_validas"], 0)
		self.assertEqual(response.data["summary"]["filas_con_error"], 1)
		self.assertTrue(
			any(
				message in response.data["rows"][0]["errors"][0]
				for message in ["no pertenece a la empresa activa", "no coincide con la empresa activa"]
			)
		)

	def test_empresa_import_unidades_uses_active_company_when_empresa_id_differs(self):
		empresa_archivo = Empresa.objects.create(
			empresa_id="EMP-ARCHIVO-UN",
			nombre="Empresa Archivo Unidades",
		)
		empresa_activa = Empresa.objects.create(
			empresa_id="EMP-ACTIVA-UN",
			nombre="Empresa Activa Unidades",
		)
		uploaded_file = SimpleUploadedFile(
			"unidades.csv",
			(
				"unidad_id,empresa_id,nombre,tipo\n"
				f"UNI-TENANT-001,{empresa_archivo.empresa_id},Unidad Tenant,Secado\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/unidades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/unidades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertIn(
			"empresa_id del archivo difiere de la empresa activa; se importara usando la empresa activa",
			preview_response.data["rows"][0].get("warnings", []),
		)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		unidad = UnidadOperativa.objects.get(unidad_id="UNI-TENANT-001")
		self.assertEqual(unidad.empresa, empresa_activa)

	def test_empresa_import_lotes_uses_active_company_when_empresa_id_differs(self):
		empresa_archivo = Empresa.objects.create(
			empresa_id="EMP-ARCHIVO-LOT",
			nombre="Empresa Archivo Lotes",
		)
		empresa_activa = Empresa.objects.create(
			empresa_id="EMP-ACTIVA-LOT",
			nombre="Empresa Activa Lotes",
		)
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa_id,empresa,fecha,especie,volumen_m3,origen,destino\n"
				f"LOTE-TENANT-001,{empresa_archivo.empresa_id},Empresa Archivo Lotes,2026-04-28,Pino radiata,10,Curico,Santiago\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertIn(
			"empresa_id del archivo difiere de la empresa activa; se importara usando la empresa activa",
			preview_response.data["rows"][0].get("warnings", []),
		)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		lote = Lote.objects.get(id_lote="LOTE-TENANT-001")
		self.assertEqual(lote.empresa, empresa_activa)

	def test_empresa_import_actividades_uses_active_company_when_empresa_id_differs(self):
		empresa_archivo = Empresa.objects.create(
			empresa_id="EMP-ARCHIVO-ACT",
			nombre="Empresa Archivo Actividades",
		)
		empresa_activa = Empresa.objects.create(
			empresa_id="EMP-ACTIVA-ACT",
			nombre="Empresa Activa Actividades",
		)
		unidad_activa = UnidadOperativa.objects.create(
			unidad_id="UNI-TENANT-ACT",
			empresa=empresa_activa,
			nombre="Unidad Activa",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			(
				"unidad_id,empresa_id,actividad,cantidad,unidad,fecha\n"
				f"{unidad_activa.unidad_id},{empresa_archivo.empresa_id},diesel,10,litros,2026-04-28\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			f"/api/empresas/{empresa_activa.empresa_id}/importaciones/actividades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["filas_validas"], 1)
		self.assertIn(
			"empresa_id del archivo difiere de la empresa activa; se importara usando la empresa activa",
			preview_response.data["rows"][0].get("warnings", []),
		)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		actividad = EmisionLote.objects.get()
		self.assertEqual(actividad.empresa, empresa_activa)
		self.assertEqual(actividad.unidad_operativa, unidad_activa)

	def test_import_lotes_deduces_empresa_from_unidad_id_without_empresa_column(self):
		"""Test that lote importer can deduce empresa from unidad_id when empresa column is missing."""
		empresa = Empresa.objects.create(
			empresa_id="EMP-LOTE-DEDUCE",
			nombre="Empresa Deduccion",
		)
		unidad = UnidadOperativa.objects.create(
			unidad_id="UNI-LOTE-DEDUCE",
			empresa=empresa,
			nombre="Unidad Deduccion",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		# CSV without "empresa" column - only id_lote, unidad_id, fecha, especie, volumen_m3, origen
		uploaded_file = SimpleUploadedFile(
			"lotes_no_empresa.csv",
			(
				"id_lote,unidad_id,fecha,especie,volumen_m3,origen\n"
				f"LOTE-DEDUCE-001,{unidad.unidad_id},2026-04-28,Pino radiata,10,Curico\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			f"/api/empresas/{empresa.empresa_id}/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			f"/api/empresas/{empresa.empresa_id}/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		lote = Lote.objects.get(id_lote="LOTE-DEDUCE-001")
		self.assertEqual(lote.empresa, empresa)
		self.assertEqual(lote.unidad_operativa, unidad)
		self.assertEqual(lote.empresa_aserradero, empresa.nombre)

	def test_dashboard_groups_by_unidad_operativa(self):
		empresa = Empresa.objects.create(empresa_id="EMP-DASH-UN", nombre="Forestal Operativa")
		aserradero = UnidadOperativa.objects.create(
			unidad_id="UNI-ASERRADERO",
			empresa=empresa,
			nombre="Aserradero Norte",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		secado = UnidadOperativa.objects.create(
			unidad_id="UNI-SECADO",
			empresa=empresa,
			nombre="Secado Sur",
			tipo=UnidadOperativa.Tipo.SECADO,
		)
		EmisionLote.objects.create(
			empresa=empresa,
			unidad_operativa=aserradero,
			actividad="Diésel - combustión móvil",
			cantidad=1,
			unidad="m3",
			factor_emision=2740,
		)
		EmisionLote.objects.create(
			empresa=empresa,
			unidad_operativa=secado,
			actividad="Electricidad Los Lagos 2023",
			cantidad=200,
			unidad="kWh",
			factor_emision=0.5,
		)

		response = self.client.get("/api/dashboard/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["emisiones_por_unidad_operativa"]["Aserradero Norte"], 2740.0)
		self.assertEqual(response.data["emisiones_por_unidad_operativa"]["Secado Sur"], 100.0)
		self.assertEqual(response.data["unidad_critica"], "Aserradero Norte")
		self.assertEqual(response.data["emisiones_por_tipo_unidad"]["Aserradero"], 2740.0)

	def test_risk_score_endpoint_returns_profile(self):
		payload = {
			"summary": {
				"total_emisiones": 350,
				"emisiones_por_actividad": {
					"diesel": 268,
					"electricidad": 82,
				},
				"emisiones_por_empresa": {
					"Empresa A": 268,
					"Empresa B": 82,
				},
				"datos": [
					{"actividad": "diesel", "emisiones": 268},
					{"actividad": "electricidad", "emisiones": 82},
				],
			}
		}

		response = self.client.post("/api/risk-score/", data=payload, format="json")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertIn("score", response.data)
		self.assertIn("label", response.data)

	def test_lotes_endpoint_creates_and_lists_lotes(self):
		payload = {
			"id_lote": "LOTE-001",
			"empresa_aserradero": "Aserradero Sur",
			"fecha": "2026-04-28",
			"especie": "Pino radiata",
			"volumen_m3": 18.75,
			"origen": "Valdivia",
			"destino": "Santiago",
		}

		create_response = self.client.post("/api/lotes/", data=payload, format="json")
		list_response = self.client.get("/api/lotes/")

		self.assertEqual(create_response.status_code, status.HTTP_201_CREATED)
		self.assertEqual(create_response.data["id_lote"], payload["id_lote"])
		self.assertAlmostEqual(float(create_response.data["masa_madera_kg"]), 9375)
		self.assertAlmostEqual(
			float(create_response.data["co2_almacenado_kg"]),
			17203.125,
		)
		self.assertEqual(Lote.objects.count(), 1)
		self.assertEqual(list_response.status_code, status.HTTP_200_OK)
		self.assertEqual(list_response.data[0]["id_lote"], payload["id_lote"])

	def test_lote_detail_endpoint_returns_lote_by_id_lote(self):
		Lote.objects.create(
			id_lote="LOTE-DET-001",
			empresa_aserradero="Aserradero Norte",
			fecha="2026-04-28",
			especie="Roble",
			volumen_m3=9.25,
			origen="Temuco",
			destino="Concepcion",
		)

		response = self.client.get("/api/lotes/LOTE-DET-001/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["id_lote"], "LOTE-DET-001")
		self.assertEqual(response.data["empresa_aserradero"], "Aserradero Norte")

	def test_especies_madera_endpoint_returns_seeded_species(self):
		response = self.client.get("/api/especies-madera/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertIn("Pino radiata", [especie["nombre"] for especie in response.data])

	def test_lote_activity_endpoint_calculates_lote_emissions(self):
		Lote.objects.create(
			id_lote="LOTE-EM-001",
			empresa_aserradero="Aserradero Centro",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=15,
			origen="Talca",
			destino="Rancagua",
		)
		payload = {
			"actividad": "diesel",
			"cantidad": 80,
			"unidad": "litros",
			"factor_emision": 2.68,
		}

		response = self.client.post(
			"/api/lotes/LOTE-EM-001/actividades/",
			data=payload,
			format="json",
		)

		self.assertEqual(response.status_code, status.HTTP_201_CREATED)
		self.assertEqual(EmisionLote.objects.count(), 1)
		self.assertEqual(response.data["actividades"][0]["actividad"], "diesel")
		self.assertEqual(response.data["actividades"][0]["emisiones_kg_co2e"], "214.400")
		self.assertEqual(response.data["total_emisiones_kg_co2e"], "214.400")

	def test_lote_carbono_endpoint_returns_net_balance(self):
		lote = Lote.objects.create(
			id_lote="LOTE-BAL-001",
			empresa_aserradero="Aserradero Balance",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/lotes/LOTE-BAL-001/carbono/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["id_lote"], "LOTE-BAL-001")
		self.assertAlmostEqual(
			float(response.data["emisiones_generadas_kg_co2e"]),
			214.4,
		)
		self.assertAlmostEqual(float(response.data["co2_almacenado_kg"]), 9175)
		self.assertAlmostEqual(
			float(response.data["balance_neto_kg_co2e"]),
			-8960.6,
		)
		self.assertEqual(response.data["estado_balance"], "favorable")
		self.assertEqual(response.data["estado_pasaporte"], "Pasaporte Verde Plus")

	def test_lote_pasaporte_endpoint_returns_mvp_classification(self):
		lote = Lote.objects.create(
			id_lote="LOTE-PASS-001",
			empresa_aserradero="Aserradero Pasaporte",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/lotes/LOTE-PASS-001/pasaporte/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["trazabilidad_score"], 100)
		self.assertEqual(response.data["completitud_score"], 100)
		self.assertEqual(response.data["factor_score"], 100)
		self.assertTrue(response.data["balance_calculado"])
		self.assertEqual(response.data["estado_pasaporte"], "Pasaporte Verde Plus")

	def test_lote_pasaporte_endpoint_blocks_lote_without_factors(self):
		Lote.objects.create(
			id_lote="LOTE-NO-PASS-001",
			empresa_aserradero="Aserradero Incompleto",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		response = self.client.get("/api/lotes/LOTE-NO-PASS-001/pasaporte/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["factor_score"], 0)
		self.assertEqual(response.data["estado_pasaporte"], "Sin pasaporte")

	def test_lote_certificado_endpoint_returns_pdf(self):
		lote = Lote.objects.create(
			id_lote="LOTE-PDF-001",
			empresa_aserradero="Aserradero PDF",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/lotes/LOTE-PDF-001/certificado/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response["Content-Type"], "application/pdf")
		self.assertIn("pasaporte-verde-LOTE-PDF-001.pdf", response["Content-Disposition"])
		self.assertTrue(response.content.startswith(b"%PDF"))

	def test_verificar_lote_endpoint_returns_public_summary(self):
		lote = Lote.objects.create(
			id_lote="LOTE-VER-001",
			empresa_aserradero="Aserradero Verificable",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/verificar/LOTE-VER-001/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["id_lote"], "LOTE-VER-001")
		self.assertEqual(response.data["estado_pasaporte"], "Pasaporte Verde Plus")
		self.assertEqual(response.data["aserradero"], "Aserradero Verificable")
		self.assertEqual(response.data["especie"], "Pino radiata")
		self.assertIn("fecha_emision", response.data)
		self.assertTrue(response.data["verificado"])

	def test_lote_documentos_endpoint_uploads_and_lists_evidence(self):
		Lote.objects.create(
			id_lote="LOTE-DOC-001",
			empresa_aserradero="Aserradero Evidencia",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			uploaded_file = SimpleUploadedFile(
				"guia.pdf",
				b"evidencia",
				content_type="application/pdf",
			)

			create_response = self.client.post(
				"/api/lotes/LOTE-DOC-001/documentos/",
				data={
					"tipo_documento": "guia_despacho",
					"fecha": "2026-04-28",
					"archivo": uploaded_file,
				},
				format="multipart",
			)
			list_response = self.client.get("/api/lotes/LOTE-DOC-001/documentos/")

			self.assertEqual(create_response.status_code, status.HTTP_201_CREATED)
			self.assertEqual(create_response.data["estado_validacion"], "pendiente")
			self.assertEqual(create_response.data["tipo_documento_label"], "Guia de despacho")
			self.assertEqual(list_response.status_code, status.HTTP_200_OK)
			self.assertEqual(len(list_response.data), 1)
			self.assertIn("archivo_url", list_response.data[0])

	def test_lote_confianza_endpoint_returns_high_confidence_with_evidence(self):
		lote = Lote.objects.create(
			id_lote="LOTE-CONF-001",
			empresa_aserradero="Aserradero Confianza",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			document = SimpleUploadedFile(
				"guia.pdf",
				b"evidencia",
				content_type="application/pdf",
			)
			self.client.post(
				"/api/lotes/LOTE-CONF-001/documentos/",
				data={
					"tipo_documento": "guia_despacho",
					"fecha": "2026-04-28",
					"archivo": document,
				},
				format="multipart",
			)
			response = self.client.get("/api/lotes/LOTE-CONF-001/confianza/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["datos_completos_score"], 100)
		self.assertEqual(response.data["factores_validos_score"], 100)
		self.assertEqual(response.data["trazabilidad_confianza_score"], 100)
		self.assertEqual(response.data["documentos_adjuntos_score"], 75)
		self.assertEqual(response.data["confianza_score"], 94)
		self.assertEqual(response.data["estado_confianza"], "Alta confianza")

	def test_lote_confianza_endpoint_returns_low_confidence_without_factors_or_docs(self):
		Lote.objects.create(
			id_lote="LOTE-CONF-LOW-001",
			empresa_aserradero="Aserradero Baja Confianza",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		response = self.client.get("/api/lotes/LOTE-CONF-LOW-001/confianza/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["documentos_adjuntos_score"], 0)
		self.assertEqual(response.data["factores_validos_score"], 0)
		self.assertEqual(response.data["estado_confianza"], "Baja confianza")

	def test_lote_transportes_endpoint_calculates_transport_emissions(self):
		Lote.objects.create(
			id_lote="LOTE-TR-001",
			empresa_aserradero="Aserradero Transporte",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros diesel",
			factor_emision=2.68,
			fuente="test",
			anio=2026,
		)
		payload = {
			"vehiculo": "Camion forestal",
			"patente": "ABCD12",
			"latitud": -33.4489,
			"longitud": -70.6693,
			"fecha_hora": "2026-04-28T12:30:00Z",
			"ruta": "Curico - Santiago",
			"distancia_km": 200,
			"consumo_estimado_litro_km": 0.3,
		}

		create_response = self.client.post(
			"/api/lotes/LOTE-TR-001/transportes/",
			data=payload,
			format="json",
		)
		lote_response = self.client.get("/api/lotes/LOTE-TR-001/")

		self.assertEqual(create_response.status_code, status.HTTP_201_CREATED)
		self.assertEqual(create_response.data["litros_calculados"], "60.000")
		self.assertEqual(
			create_response.data["emisiones_transporte_kg_co2e"],
			"160.800",
		)
		self.assertEqual(lote_response.data["transportes"][0]["patente"], "ABCD12")
		self.assertEqual(lote_response.data["actividades"][0]["actividad"], "transporte")
		self.assertEqual(lote_response.data["emisiones_kg_co2e"], "160.800")

	def test_factores_emision_endpoint_returns_selectable_labels(self):
		FactorEmision.objects.create(
			actividad="Diesel - combustion movil",
			unidad="litros",
			factor_emision=2.68,
			fuente="GHG",
			anio=2026,
		)

		response = self.client.get("/api/factores-emision/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data[0]["actividad"], "Diesel - combustion movil")
		self.assertIn("Diesel - combustion movil", response.data[0]["label"])

	def test_documento_ocr_suggests_data_without_applying_calculation(self):
		Lote.objects.create(
			id_lote="LOTE-OCR-001",
			empresa_aserradero="Aserradero OCR",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			document = SimpleUploadedFile(
				"factura.txt",
				(
					b"fecha: 2026-04-28\nproveedor: Copec\n"
					b"litros: 80\npatente: ABCD12\nmonto: 120000\n"
				),
				content_type="text/plain",
			)
			document_response = self.client.post(
				"/api/lotes/LOTE-OCR-001/documentos/",
				data={
					"tipo_documento": "factura_combustible",
					"fecha": "2026-04-28",
					"archivo": document,
				},
				format="multipart",
			)
			ocr_response = self.client.post(
				f"/api/documentos/{document_response.data['id']}/ocr/",
			)
			lote_response = self.client.get("/api/lotes/LOTE-OCR-001/")

		self.assertEqual(ocr_response.status_code, status.HTTP_201_CREATED)
		self.assertEqual(ocr_response.data["estado_revision"], "pendiente")
		self.assertEqual(ocr_response.data["datos_sugeridos"]["litros_combustible"], 80.0)
		self.assertEqual(len(lote_response.data["actividades"]), 0)

	def test_extraccion_validar_applies_human_reviewed_data_to_calculation(self):
		Lote.objects.create(
			id_lote="LOTE-OCR-VAL-001",
			empresa_aserradero="Aserradero OCR",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			document = SimpleUploadedFile(
				"factura.txt",
				b"litros: 80\nkWh: 120\n",
				content_type="text/plain",
			)
			document_response = self.client.post(
				"/api/lotes/LOTE-OCR-VAL-001/documentos/",
				data={
					"tipo_documento": "factura_combustible",
					"fecha": "2026-04-28",
					"archivo": document,
				},
				format="multipart",
			)
			ocr_response = self.client.post(
				f"/api/documentos/{document_response.data['id']}/ocr/",
			)
			self.assertEqual(ocr_response.data["datos_sugeridos"]["kwh"], 120.0)
			validate_response = self.client.post(
				f"/api/extracciones/{ocr_response.data['id']}/validar/",
				data={
					"datos_validados": {
						"litros_combustible": 80,
						"kwh": 120,
					},
					"aplicar_calculo": True,
				},
				format="json",
			)

		self.assertEqual(validate_response.status_code, status.HTTP_200_OK)
		self.assertEqual(validate_response.data["actividades_creadas"], 2)
		self.assertEqual(
			validate_response.data["extraccion"]["estado_revision"],
			"validado",
		)
		self.assertEqual(validate_response.data["lote"]["emisiones_kg_co2e"], "262.400")

	@patch("apps.analytics.services.rutas.osrm_route_distance_km")
	@patch("apps.analytics.services.rutas.geocode_location")
	def test_calcular_distancia_ruta_endpoint_returns_distance(self, mock_geocode, mock_osrm):
		mock_geocode.side_effect = [
			{"lat": -33.4489, "lon": -70.6693, "display_name": "Santiago"},
			{"lat": -34.1708, "lon": -70.7444, "display_name": "Curico"},
		]
		mock_osrm.return_value = None

		response = self.client.post(
			"/api/rutas/calcular-distancia/",
			data={"origen": "Santiago", "destino": "Curico"},
			format="json",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["metodo"], "geodesic_fallback")
		self.assertGreater(response.data["distancia_km"], 0)
		self.assertEqual(len(response.data["route_geometry"]), 2)

	@patch("apps.analytics.services.rutas._http_json")
	def test_calcular_distancia_ruta_returns_osrm_geometry_for_leaflet(self, mock_http_json):
		mock_http_json.return_value = {
			"code": "Ok",
			"routes": [
				{
					"distance": 12345.6,
					"duration": 1800,
					"geometry": {
						"coordinates": [
							[-70.6693, -33.4489],
							[-70.7, -33.7],
							[-70.7444, -34.1708],
						]
					},
				}
			],
		}

		response = self.client.post(
			"/api/rutas/calcular-distancia/",
			data={
				"origen": "Santiago",
				"destino": "Curico",
				"origen_coords": {"lat": -33.4489, "lon": -70.6693},
				"destino_coords": {"lat": -34.1708, "lon": -70.7444},
			},
			format="json",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["metodo"], "osrm_driving")
		self.assertEqual(response.data["distancia_km"], 12.346)
		self.assertEqual(response.data["duracion_min"], 30)
		self.assertEqual(
			response.data["route_geometry"],
			[
				[-33.4489, -70.6693],
				[-33.7, -70.7],
				[-34.1708, -70.7444],
			],
		)

	@unittest.skipIf(PdfReader is None, "pypdf no esta instalado")
	def test_documento_ocr_extracts_text_from_pdf(self):
		Lote.objects.create(
			id_lote="LOTE-PDF-001",
			empresa_aserradero="Aserradero PDF",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			buffer = BytesIO()
			pdf = canvas.Canvas(buffer)
			pdf.drawString(72, 750, "fecha: 2026-04-28")
			pdf.drawString(72, 730, "proveedor: Copec")
			pdf.drawString(72, 710, "litros: 120")
			pdf.drawString(72, 690, "patente: ABCD12")
			pdf.save()
			uploaded = SimpleUploadedFile(
				"factura.pdf",
				buffer.getvalue(),
				content_type="application/pdf",
			)
			document_response = self.client.post(
				"/api/lotes/LOTE-PDF-001/documentos/",
				data={
					"tipo_documento": "factura_combustible",
					"fecha": "2026-04-28",
					"archivo": uploaded,
				},
				format="multipart",
			)
			ocr_response = self.client.post(
				f"/api/documentos/{document_response.data['id']}/ocr/",
			)

		self.assertEqual(ocr_response.status_code, status.HTTP_201_CREATED)
		self.assertIn("fecha: 2026-04-28", ocr_response.data["texto_extraido"])
		self.assertEqual(ocr_response.data["datos_sugeridos"]["litros_combustible"], 120.0)

	@unittest.skipIf(Document is None, "python-docx no esta instalado")
	def test_documento_extract_json_returns_structured_payload_from_docx(self):
		Lote.objects.create(
			id_lote="LOTE-DOCX-001",
			empresa_aserradero="Aserradero DOCX",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)

		with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
			buffer = BytesIO()
			document = Document()
			document.add_paragraph("factura combustible")
			document.add_paragraph("fecha: 2026-04-28")
			document.add_paragraph("litros diesel: 120")
			document.add_paragraph("patente: ABCD-12")
			document.add_paragraph("id_lote: LOTE-001")
			document.save(buffer)
			uploaded = SimpleUploadedFile(
				"factura.docx",
				buffer.getvalue(),
				content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
			)
			document_response = self.client.post(
				"/api/lotes/LOTE-DOCX-001/documentos/",
				data={
					"tipo_documento": "factura_combustible",
					"fecha": "2026-04-28",
					"archivo": uploaded,
				},
				format="multipart",
			)
			structured_response = self.client.post(
				f"/api/documentos/{document_response.data['id']}/extract-json/",
			)

		self.assertEqual(structured_response.status_code, status.HTTP_200_OK)
		self.assertEqual(structured_response.data["tipo_documento"], "factura_combustible")
		self.assertEqual(structured_response.data["id_lote"], "LOTE-001")
		self.assertEqual(structured_response.data["litros_diesel"], 120.0)

	def test_integracion_lote_endpoint_returns_bim_payload(self):
		lote = Lote.objects.create(
			id_lote="LOTE-BIM-001",
			empresa_aserradero="Aserradero BIM",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=12,
			origen="Curico",
			destino="Santiago",
		)
		lote.actividades.create(
			actividad="diesel",
			cantidad=80,
			unidad="litros",
			factor_emision=2.68,
		)

		response = self.client.get("/api/integraciones/lotes/LOTE-BIM-001/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["lote"], "LOTE-BIM-001")
		self.assertEqual(response.data["producto"], "Pino radiata dimensionado")
		self.assertEqual(response.data["volumen_m3"], 12.0)
		self.assertAlmostEqual(response.data["emisiones_kgco2e"], 214.4)
		self.assertIn("bim", response.data)
		self.assertEqual(
			response.data["bim"]["property_set"],
			"Pset_HuellaPasaporteVerde",
		)

	def test_integracion_lote_exports_json_and_csv(self):
		Lote.objects.create(
			id_lote="LOTE-EXPORT-001",
			empresa_aserradero="Aserradero Export",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=12,
			origen="Curico",
			destino="Santiago",
		)

		json_response = self.client.get(
			"/api/integraciones/lotes/LOTE-EXPORT-001/export.json"
		)
		csv_response = self.client.get(
			"/api/integraciones/lotes/LOTE-EXPORT-001/export.csv"
		)

		self.assertEqual(json_response.status_code, status.HTTP_200_OK)
		self.assertEqual(json_response["Content-Type"], "application/json")
		self.assertIn("LOTE-EXPORT-001-bim.json", json_response["Content-Disposition"])
		self.assertEqual(csv_response.status_code, status.HTTP_200_OK)
		self.assertEqual(csv_response["Content-Type"], "text/csv")
		self.assertIn("lote,producto", csv_response.content.decode("utf-8"))

	def test_integracion_lote_ficha_tecnica_returns_material_sheet(self):
		Lote.objects.create(
			id_lote="LOTE-FICHA-001",
			empresa_aserradero="Aserradero Ficha",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=12,
			origen="Curico",
			destino="Santiago",
		)

		response = self.client.get(
			"/api/integraciones/lotes/LOTE-FICHA-001/ficha-tecnica/"
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["lote"], "LOTE-FICHA-001")
		self.assertIn("densidad_kg_m3", response.data["ficha_tecnica"])
		self.assertIn("properties", response.data["bim"])

	def test_import_factores_preview_and_confirm_persists_valid_rows(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"diesel,litros,2.68,DEFRA,2025\n"
				"electricidad,kWh,,SEC,2025\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["total_filas"], 2)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertEqual(preview_response.data["summary"]["con_error"], 1)
		self.assertEqual(FactorEmision.objects.count(), 0)

		confirm_response = self.client.post(
			"/api/importaciones/factores/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["creados"], 1)
		self.assertEqual(FactorEmision.objects.count(), 1)
		factor = FactorEmision.objects.get()
		self.assertEqual(factor.categoria, "Combustible")
		self.assertEqual(factor.actividad_key, "diesel")

	def test_factor_classifier_normalizes_keys_and_infers_categories(self):
		self.assertEqual(normalize_key("Diésel - combustión móvil"), "diesel_combustion_movil")
		self.assertEqual(normalize_key("Cartón virgen"), "carton_virgen")
		self.assertEqual(
			normalize_key("Residuos de cartón - Reciclaje"),
			"residuos_carton_reciclaje",
		)
		cases = [
			("Diésel - combustión móvil", "m3", "Combustible"),
			("diesel_movil", "m3", "Combustible"),
			("Camión diésel rígido promedio", "t-km", "Transporte"),
			("barco de carga promedio", "t-km", "Transporte"),
			("Electricidad Los Lagos 2023", "kWh", "Electricidad"),
			("electricidad_sen_2023", "kWh", "Electricidad"),
			("Suministro de agua potable", "m3", "Agua"),
			("Tratamiento de agua", "m3", "Agua"),
			("Cartón virgen", "t", "Materiales"),
			("carton_virgen", "t", "Materiales"),
			("Residuos de cartón - Reciclaje", "t", "Residuos"),
			("residuos_carton_reciclaje", "t", "Residuos"),
			("Refrigerante R507", "kg", "Refrigerantes"),
		]
		for actividad, unidad, expected in cases:
			with self.subTest(actividad=actividad):
				self.assertEqual(infer_categoria(actividad, unidad), expected)

	def test_activity_semantics_detects_extended_internal_names(self):
		self.assertTrue(is_diesel_activity({"actividad": "Diésel - combustión móvil"}))
		self.assertTrue(is_diesel_activity({"actividad": "diesel_movil"}))
		self.assertTrue(
			is_electricity_activity({"actividad": "Electricidad Los Lagos 2023"})
		)

	def test_decision_engine_simulates_extended_diesel_activity(self):
		rows = [
			{
				"empresa": "Aserradero Verde Sur",
				"actividad": "Diésel - combustión móvil",
				"actividad_key": "diesel_combustion_movil",
				"cantidad": 1,
				"unidad": "m3",
				"factor_emision": 2740,
				"emisiones": 2740,
			}
		]

		simulated = simulate_rows(rows, diesel_reduction=50)

		self.assertEqual(simulated[0]["cantidad"], 0.5)
		self.assertEqual(simulated[0]["emisiones"], 1370)

	def test_decision_engine_optimizer_reduces_extended_diesel_dataset(self):
		rows = [
			{
				"empresa": "Aserradero Verde Sur",
				"actividad": "Diésel - combustión móvil",
				"actividad_key": "diesel_combustion_movil",
				"cantidad": 1,
				"unidad": "m3",
				"factor_emision": 2740,
				"emisiones": 2740,
			},
			{
				"empresa": "Aserradero Verde Sur",
				"actividad": "Electricidad Los Lagos 2023",
				"actividad_key": "electricidad_los_lagos_2023",
				"cantidad": 100,
				"unidad": "kWh",
				"factor_emision": 0.6242,
				"emisiones": 62.42,
			},
		]

		optimized = optimize_rows(rows)
		risk = calculate_risk_profile(
			{
				"total_emisiones": 2802.42,
				"emisiones_por_actividad": {
					"Diésel - combustión móvil": 2740,
					"Electricidad Los Lagos 2023": 62.42,
				},
				"emisiones_por_empresa": {"Aserradero Verde Sur": 2802.42},
				"datos": rows,
			},
			optimized,
		)

		self.assertGreater(optimized["evaluatedScenarios"], 1)
		self.assertGreater(optimized["dieselReduction"], 0)
		self.assertGreater(optimized["reductionPct"], 0)
		self.assertLess(optimized["simulatedTotal"], optimized["currentTotal"])
		self.assertTrue(risk["factors"]["dieselPresent"])

	def test_import_factores_preview_includes_classification_fields(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"Camión diésel rígido promedio,t-km,0.1782,DEFRA,2025\n"
				"Factor desconocido,kg,1.2,Manual,2025\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["rows"][0]["data"]["categoria"], "Transporte")
		self.assertEqual(
			response.data["rows"][0]["data"]["actividad_key"],
			"camion_diesel_rigido_promedio",
		)
		self.assertEqual(response.data["rows"][1]["data"]["categoria"], "Otros")
		self.assertIn("Categoria no detectada", response.data["rows"][1]["warnings"][0])

	def test_import_factores_with_valid_category_respects_file_value(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio,categoria\n"
				"Diésel - combustión móvil,m3,2740,Manual,2025,Transporte\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/factores/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(FactorEmision.objects.get().categoria, "Transporte")

	def test_import_factores_with_otros_category_still_infers_when_clear(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio,categoria\n"
				"Diésel - combustión móvil,m3,2740,Manual,2025,Otros\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["rows"][0]["data"]["categoria"], "Combustible")

	def test_import_factores_invalid_category_becomes_otros_with_warning(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio,categoria\n"
				"Cartón virgen,t,910.48,Manual,2025,Categoria inventada\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["rows"][0]["data"]["categoria"], "Otros")
		self.assertIn("Categoria invalida", response.data["rows"][0]["warnings"][0])

	def test_import_factores_avoids_duplicates_by_activity_key(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"Cartón virgen,t,910.48,Manual,2025\n"
				"carton_virgen,t,910.48,Manual,2025\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["duplicadas"], 1)

	def test_import_factores_treats_equivalent_readable_and_slug_rows_as_duplicate(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"barco_carga_promedio,t-km,0.0161,Fuente antigua,2024\n"
				"barco de carga promedio,t-km,0.0161,Fuente nueva,2024\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/factores/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.data["summary"]["duplicadas"], 1)
		self.assertEqual(FactorEmision.objects.count(), 1)
		factor = FactorEmision.objects.get()
		self.assertEqual(factor.actividad, "barco de carga promedio")
		self.assertEqual(confirm_response.data["creados"], 1)

	def test_limpiar_factores_duplicados_keeps_readable_row(self):
		FactorEmision.objects.create(
			actividad="camion_diesel_rigido_promedio",
			unidad="t-km",
			factor_emision=0.1782,
			fuente="HuellaChile Nivel básico v3 2024 / DEFRA 2023",
			anio=2024,
		)
		FactorEmision.objects.create(
			actividad="camión diésel rígido promedio",
			unidad="t-km",
			factor_emision=0.1782,
			fuente="HuellaChile 2024 / DEFRA 2023",
			anio=2024,
		)

		call_command("limpiar_factores_duplicados")

		self.assertEqual(FactorEmision.objects.count(), 1)
		factor = FactorEmision.objects.get()
		self.assertEqual(factor.actividad, "camión diésel rígido promedio")
		self.assertEqual(factor.actividad_key, "camion_diesel_rigido_promedio")

	def test_factores_catalogo_returns_grouped_categories(self):
		FactorEmision.objects.create(
			actividad="Electricidad Los Lagos 2023",
			unidad="kWh",
			factor_emision=0.6242,
			fuente="SEC",
			anio=2023,
		)

		response = self.client.get("/api/factores/catalogo/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["Electricidad"][0]["actividad_key"], "electricidad_los_lagos_2023")

	def test_import_factores_preview_accepts_xlsx_and_detects_duplicates(self):
		workbook = Workbook()
		sheet = workbook.active
		sheet.append(["actividad", "unidad", "factor_emision", "fuente", "anio"])
		sheet.append(["Diesel  ", " Litros ", 2.68, " DEFRA ", 2025])
		sheet.append(["diesel", "litros", 2.68, "DEFRA", 2025])
		buffer = BytesIO()
		workbook.save(buffer)
		workbook.close()
		uploaded_file = SimpleUploadedFile(
			"factores.xlsx",
			buffer.getvalue(),
			content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["total_filas"], 2)
		self.assertEqual(preview_response.data["summary"]["duplicadas"], 1)
		self.assertEqual(preview_response.data["summary"]["validas"], 1)
		self.assertEqual(preview_response.data["rows"][0]["data"]["actividad"], "Diesel")
		self.assertEqual(preview_response.data["rows"][0]["data"]["unidad"], "Litros")

	def test_import_factores_preview_rejects_missing_columns(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			("actividad,unidad,factor_emision,anio\n" "diesel,litros,2.68,2025\n").encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
		self.assertIn("Faltan columnas obligatorias", response.data["error"])

	def test_import_factores_preview_marks_invalid_factor(self):
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"diesel,litros,-2.68,DEFRA,2025\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["validas"], 0)
		self.assertEqual(response.data["summary"]["con_error"], 1)
		self.assertTrue(response.data["rows"][0]["errors"])

	def test_import_factores_confirm_updates_existing_factor(self):
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.0,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"factores.csv",
			(
				"actividad,unidad,factor_emision,fuente,anio\n"
				"Diesel, Litros ,2.68, DEFRA ,2025\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/factores/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/factores/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		factor = FactorEmision.objects.get()
		self.assertEqual(preview_response.data["summary"]["posibles_actualizaciones"], 1)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["actualizados"], 1)
		self.assertEqual(str(factor.factor_emision), "2.680000")

	def test_import_empresa_completa_accepts_ui_headers_and_file_factors(self):
		workbook = Workbook()
		empresa_sheet = workbook.active
		empresa_sheet.title = "empresa"
		empresa_sheet.append([
			"ID Empresa",
			"Nombre",
			"RUT",
			"Región",
			"Comuna",
			"Dirección",
			"Rubro",
			"Email",
			"Teléfono",
			"Contacto",
			"Observaciones",
		])
		empresa_sheet.append([
			"EMP-COMPLETA",
			"Empresa Completa Demo",
			"76.543.210-9",
			"Biobío",
			"Concepción",
			"Ruta 5 Sur",
			"Forestal",
			"demo@empresa.cl",
			"+56 9 1111 2222",
			"Contacto Demo",
			"Importación completa de prueba",
		])

		factores_sheet = workbook.create_sheet("factores")
		factores_sheet.append(["Actividad", "Unidad", "Factor de Emisión", "Fuente", "Año"])
		factores_sheet.append(["Diesel", "litros", 2.68, "Demo", 2025])

		unidades_sheet = workbook.create_sheet("unidades")
		unidades_sheet.append(["ID Unidad", "Nombre", "Tipo", "Región", "Comuna", "Dirección"])
		unidades_sheet.append(["UNI-COMPLETA", "Aserradero", "Aserradero", "Biobío", "Los Ángeles", "Ruta 5 Sur"])

		lotes_sheet = workbook.create_sheet("lotes")
		lotes_sheet.append(["ID Lote", "ID Unidad", "Fecha", "Especie", "Volumen (m³)", "Origen"])
		lotes_sheet.append(["LOTE-COMPLETA", "UNI-COMPLETA", "2025-01-15", "Pino radiata", 12.5, "Predio Demo"])

		actividades_sheet = workbook.create_sheet("actividades")
		actividades_sheet.append(["ID Unidad", "ID Lote", "Actividad", "Cantidad", "Unidad", "Fecha"])
		actividades_sheet.append(["UNI-COMPLETA", "LOTE-COMPLETA", "Diesel", 10, "litros", "2025-01-15"])

		buffer = BytesIO()
		workbook.save(buffer)
		workbook.close()
		uploaded_file = SimpleUploadedFile(
			"empresa_completa.xlsx",
			buffer.getvalue(),
			content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
		)

		preview_response = self.client.post(
			"/api/importaciones/empresa-completa/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/empresa-completa/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["blocking_errors"], [])
		self.assertEqual(preview_response.data["factores"]["validos"], 1)
		self.assertEqual(preview_response.data["actividades"]["validas"], 1)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["creados"], 1)
		self.assertEqual(confirm_response.data["factores_creados"], 1)
		self.assertEqual(Empresa.objects.get(empresa_id="EMP-COMPLETA").nombre, "Empresa Completa Demo")
		self.assertEqual(UnidadOperativa.objects.get(unidad_id="UNI-COMPLETA").empresa.empresa_id, "EMP-COMPLETA")
		self.assertEqual(Lote.objects.get(id_lote="LOTE-COMPLETA").unidad_operativa.unidad_id, "UNI-COMPLETA")
		self.assertEqual(Lote.objects.get(id_lote="LOTE-COMPLETA").empresa_aserradero, "Empresa Completa Demo")
		self.assertEqual(EmisionLote.objects.count(), 1)
		self.assertAlmostEqual(float(EmisionLote.objects.get().emisiones_kg_co2e), 26.8)

	def test_import_empresa_completa_accepts_unit_names_in_lotes_and_activities(self):
		workbook = Workbook()
		empresa_sheet = workbook.active
		empresa_sheet.title = "empresa"
		empresa_sheet.append([
			"ID Empresa",
			"Nombre",
			"RUT",
			"Región",
			"Comuna",
			"Dirección",
			"Rubro",
			"Email",
			"Teléfono",
			"Contacto",
			"Observaciones",
		])
		empresa_sheet.append([
			"EMP-NOMBRES",
			"Empresa Nombres Demo",
			"76.999.888-7",
			"Biobío",
			"Concepción",
			"Ruta 5",
			"Forestal",
			"demo@nombres.cl",
			"+56 9 2222 3333",
			"Contacto",
			"Prueba de nombres",
		])

		factores_sheet = workbook.create_sheet("factores")
		factores_sheet.append(["Actividad", "Unidad", "Factor de Emisión", "Fuente", "Año"])
		factores_sheet.append(["Diesel", "litros", 2.68, "Demo", 2025])

		unidades_sheet = workbook.create_sheet("unidades")
		unidades_sheet.append(["ID Unidad", "Nombre", "Tipo", "Región", "Comuna", "Dirección"])
		unidades_sheet.append(["UNI-NOMBRE-01", "Aserradero Nombres", "Aserradero", "Biobío", "Los Ángeles", "Ruta 5"])

		lotes_sheet = workbook.create_sheet("lotes")
		lotes_sheet.append(["ID Lote", "ID Unidad", "Fecha", "Especie", "Volumen (m³)", "Origen"])
		lotes_sheet.append(["LOTE-NOMBRE-01", "Aserradero Nombres", "2025-01-15", "Pino radiata", 12.5, "Predio Demo"])

		actividades_sheet = workbook.create_sheet("actividades")
		actividades_sheet.append(["ID Unidad", "ID Lote", "Actividad", "Cantidad", "Unidad", "Fecha"])
		actividades_sheet.append(["Aserradero Nombres", "LOTE-NOMBRE-01", "Diesel", 10, "litros", "2025-01-15"])

		buffer = BytesIO()
		workbook.save(buffer)
		workbook.close()
		uploaded_file = SimpleUploadedFile(
			"empresa_completa_nombres.xlsx",
			buffer.getvalue(),
			content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
		)

		preview_response = self.client.post(
			"/api/importaciones/empresa-completa/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/empresa-completa/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["blocking_errors"], [])
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["lotes_creados"], 1)
		self.assertEqual(confirm_response.data["actividades_creadas"], 1)
		self.assertEqual(Lote.objects.get(id_lote="LOTE-NOMBRE-01").unidad_operativa.nombre, "Aserradero Nombres")
		self.assertEqual(EmisionLote.objects.get().unidad_operativa.nombre, "Aserradero Nombres")

	def test_import_actividades_preview_detects_missing_factor_and_lote(self):
		Lote.objects.create(
			id_lote="LOTE-IMP-001",
			empresa_aserradero="Aserradero Import",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			(
				"id_lote,actividad,cantidad,unidad,fecha\n"
				"LOTE-IMP-001,diesel,80,litros,2026-04-28\n"
				"LOTE-NUEVO,diesel,20,litros,2026-04-28\n"
				"LOTE-IMP-001,gas,10,m3,2026-04-28\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		valid_rows = [
			row for row in preview_response.data["rows"] if row["status"] == "valid"
		]
		confirm_response = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"rows": valid_rows},
			format="json",
		)

		self.assertEqual(preview_response.status_code, status.HTTP_200_OK)
		self.assertEqual(preview_response.data["summary"]["filas_validas"], 1)
		self.assertEqual(preview_response.data["summary"]["filas_con_error"], 2)
		self.assertEqual(preview_response.data["summary"]["factores_faltantes"], 1)
		self.assertEqual(preview_response.data["summary"]["lotes_nuevos_detectados"], 1)
		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["created"], 1)
		self.assertAlmostEqual(float(EmisionLote.objects.get().emisiones_kg_co2e), 214.4)

	def test_import_unidades_validates_empresa_id(self):
		Empresa.objects.create(empresa_id="EMP-IMPORT", nombre="Empresa Importadora")
		uploaded_file = SimpleUploadedFile(
			"unidades.csv",
			(
				"unidad_id,empresa_id,nombre,tipo\n"
				"UNI-OK,EMP-IMPORT,Secado,Secado\n"
				"UNI-BAD,EMP-NO-EXISTE,Despacho,Despacho\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/unidades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["validas"], 1)
		self.assertEqual(response.data["summary"]["con_error"], 1)
		self.assertIn("empresa_id no existe", response.data["rows"][1]["errors"][0])

	def test_import_actividades_inherits_empresa_and_unidad_from_lote(self):
		empresa = Empresa.objects.create(empresa_id="EMP-ACT-IMP", nombre="Empresa Actividades")
		unidad = UnidadOperativa.objects.create(
			unidad_id="UNI-ACT-IMP",
			empresa=empresa,
			nombre="Aserradero Principal",
			tipo=UnidadOperativa.Tipo.ASERRADERO,
		)
		Lote.objects.create(
			id_lote="LOTE-ACT-IMP",
			empresa=empresa,
			unidad_operativa=unidad,
			empresa_aserradero=empresa.nombre,
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			(
				"id_lote,actividad,cantidad,unidad,fecha\n"
				"LOTE-ACT-IMP,diesel,10,litros,2026-04-28\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)
		actividad = EmisionLote.objects.get()

		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(actividad.empresa, empresa)
		self.assertEqual(actividad.unidad_operativa, unidad)
		self.assertEqual(actividad.tipo_asignacion, EmisionLote.TipoAsignacion.LOTE)

	def test_import_actividades_once_creates_21_records_and_dashboard_uses_21(self):
		Lote.objects.create(
			id_lote="LOTE-ACT-21",
			empresa_aserradero="Aserradero Demo",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2,
			fuente="Demo",
			anio=2026,
		)
		rows = [
			f"LOTE-ACT-21,diesel,{index},litros,2026-04-{index:02d}"
			for index in range(1, 22)
		]
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			("id_lote,actividad,cantidad,unidad,fecha\n" + "\n".join(rows) + "\n").encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)
		dashboard_response = self.client.get("/api/dashboard/")

		self.assertEqual(preview_response.data["summary"]["filas_validas"], 21)
		self.assertEqual(confirm_response.data["creados"], 21)
		self.assertEqual(EmisionLote.objects.count(), 21)
		self.assertEqual(dashboard_response.data["cantidad_registros"], 21)
		self.assertAlmostEqual(dashboard_response.data["total_emisiones"], 462.0)

	def test_import_actividades_confirm_same_batch_twice_does_not_duplicate(self):
		Lote.objects.create(
			id_lote="LOTE-BATCH-001",
			empresa_aserradero="Aserradero Batch",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		uploaded_file = SimpleUploadedFile(
			"actividades.csv",
			(
				"id_lote,actividad,cantidad,unidad,fecha\n"
				"LOTE-BATCH-001,diesel,80,litros,2026-04-28\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		preview_response = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		first_confirm = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)
		second_confirm = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(first_confirm.data["creados"], 1)
		self.assertEqual(second_confirm.data["creados"], 0)
		self.assertEqual(second_confirm.data["message"], "Esta importación ya fue aplicada.")
		self.assertEqual(EmisionLote.objects.count(), 1)

	def test_import_actividades_same_file_twice_detects_db_duplicate_and_omits(self):
		Lote.objects.create(
			id_lote="LOTE-DUP-ACT",
			empresa_aserradero="Aserradero Dup",
			fecha="2026-04-28",
			especie="Pino radiata",
			volumen_m3=10,
			origen="Curico",
			destino="Santiago",
		)
		FactorEmision.objects.create(
			actividad="diesel",
			unidad="litros",
			factor_emision=2.68,
			fuente="DEFRA",
			anio=2025,
		)
		content = (
			"id_lote,actividad,cantidad,unidad,fecha\n"
			"LOTE-DUP-ACT,diesel,80,litros,2026-04-28\n"
		).encode("utf-8")

		first_preview = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": SimpleUploadedFile("actividades.csv", content, content_type="text/csv")},
			format="multipart",
		)
		self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": first_preview.data["batch_id"]},
			format="json",
		)
		second_preview = self.client.post(
			"/api/importaciones/actividades/preview/",
			data={"file": SimpleUploadedFile("actividades.csv", content, content_type="text/csv")},
			format="multipart",
		)
		second_confirm = self.client.post(
			"/api/importaciones/actividades/confirm/",
			data={"batch_id": second_preview.data["batch_id"]},
			format="json",
		)

		self.assertEqual(second_preview.data["summary"]["filas_validas"], 0)
		self.assertEqual(second_preview.data["summary"]["duplicados"], 1)
		self.assertEqual(second_confirm.data["creados"], 0)
		self.assertEqual(second_confirm.data["duplicados"], 1)
		self.assertEqual(EmisionLote.objects.count(), 1)

	def make_lotes_xlsx(self, rows):
		workbook = Workbook()
		sheet = workbook.active
		sheet.append(
			[
				"id_lote",
				"empresa",
				"fecha",
				"especie",
				"volumen_m3",
				"origen",
				"destino",
				"densidad_kg_m3",
				"porcentaje_carbono",
			]
		)
		for row in rows:
			sheet.append(row)
		buffer = BytesIO()
		workbook.save(buffer)
		buffer.seek(0)
		return SimpleUploadedFile(
			"lotes.xlsx",
			buffer.read(),
			content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
		)

	def test_import_lotes_preview_accepts_valid_csv_without_writing(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino\n"
				" lote-001 , Aserradero Sur ,2026-04-28, Pino radiata ,10,Curico,Santiago\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["total_filas"], 1)
		self.assertEqual(response.data["summary"]["validas"], 1)
		self.assertEqual(response.data["summary"]["posibles_creaciones"], 1)
		self.assertEqual(response.data["rows"][0]["data"]["id_lote"], "LOTE-001")
		self.assertEqual(response.data["rows"][0]["data"]["especie"], "pino radiata")
		self.assertEqual(Lote.objects.count(), 0)

	def test_import_lotes_preview_accepts_valid_xlsx(self):
		uploaded_file = self.make_lotes_xlsx(
			[["LOTE-XLSX-001", "Aserradero XLSX", "2026-04-28", "Roble", 8, "Temuco", "Concepcion", None, None]]
		)

		response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["validas"], 1)
		self.assertEqual(response.data["rows"][0]["data"]["id_lote"], "LOTE-XLSX-001")

	def test_import_lotes_preview_rejects_missing_columns(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			"id_lote,empresa,fecha\nLOTE-001,Aserradero,2026-04-28\n".encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
		self.assertIn("Faltan columnas obligatorias", response.data["error"])

	def test_import_lotes_preview_marks_invalid_volume_and_date(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino\n"
				"LOTE-BAD,Aserradero,fecha-mala,Pino radiata,-1,Curico,Santiago\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["con_error"], 1)
		self.assertIn("fecha no tiene un formato valido", response.data["rows"][0]["errors"])
		self.assertIn("volumen_m3 debe ser mayor que cero", response.data["rows"][0]["errors"])

	def test_import_lotes_preview_marks_duplicate_id_in_file(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino\n"
				"LOTE-DUP,Aserradero,2026-04-28,Pino radiata,10,Curico,Santiago\n"
				"LOTE-DUP,Aserradero,2026-04-28,Pino radiata,12,Curico,Santiago\n"
			).encode("utf-8"),
			content_type="text/csv",
		)

		response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertEqual(response.data["summary"]["duplicadas"], 1)
		self.assertEqual(response.data["rows"][1]["status"], "error")

	def test_import_lotes_confirm_creates_lotes_and_audit_event(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino,tipo_producto,estado,observaciones\n"
				"LOTE-CREATE,Aserradero Nuevo,2026-04-28,Pino radiata,10,Curico,Santiago,tabla,pendiente,importado\n"
			).encode("utf-8"),
			content_type="text/csv",
		)
		preview_response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)

		confirm_response = self.client.post(
			"/api/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(confirm_response.status_code, status.HTTP_200_OK)
		self.assertEqual(confirm_response.data["creados"], 1)
		lote = Lote.objects.get(id_lote="LOTE-CREATE")
		self.assertEqual(lote.empresa_aserradero, "Aserradero Nuevo")
		self.assertEqual(lote.tipo_producto, "tabla")
		self.assertEqual(lote.historial_cambios.first().fuente, "importador_lotes")
		self.assertEqual(lote.historial_cambios.first().tipo, HistorialCambioLote.TipoCambio.IMPORTADO)

	def test_import_lotes_confirm_updates_existing_lote(self):
		Lote.objects.create(
			id_lote="LOTE-UPD",
			empresa_aserradero="Empresa Antigua",
			fecha="2026-04-01",
			especie="Pino radiata",
			volumen_m3=5,
			origen="Talca",
			destino="Rancagua",
		)
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino\n"
				"LOTE-UPD,Empresa Nueva,2026-04-28,Roble,12,Temuco,Concepcion\n"
			).encode("utf-8"),
			content_type="text/csv",
		)
		preview_response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		confirm_response = self.client.post(
			"/api/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		self.assertEqual(preview_response.data["summary"]["posibles_actualizaciones"], 1)
		self.assertEqual(confirm_response.data["actualizados"], 1)
		lote = Lote.objects.get(id_lote="LOTE-UPD")
		self.assertEqual(lote.empresa_aserradero, "Empresa Nueva")
		self.assertEqual(str(lote.volumen_m3), "12.000")

	def test_import_lotes_carbono_uses_imported_density_when_species_unknown(self):
		uploaded_file = SimpleUploadedFile(
			"lotes.csv",
			(
				"id_lote,empresa,fecha,especie,volumen_m3,origen,destino,densidad_kg_m3,porcentaje_carbono\n"
				"LOTE-CARB,Aserradero Carbono,2026-04-28,especie propia,10,Curico,Santiago,500,0.5\n"
			).encode("utf-8"),
			content_type="text/csv",
		)
		preview_response = self.client.post(
			"/api/importaciones/lotes/preview/",
			data={"file": uploaded_file},
			format="multipart",
		)
		self.client.post(
			"/api/importaciones/lotes/confirm/",
			data={"batch_id": preview_response.data["batch_id"]},
			format="json",
		)

		response = self.client.get("/api/lotes/LOTE-CARB/carbono/")

		self.assertEqual(response.status_code, status.HTTP_200_OK)
		self.assertAlmostEqual(float(response.data["masa_madera_kg"]), 5000)
		self.assertAlmostEqual(float(response.data["co2_almacenado_kg"]), 9175)
		self.assertTrue(response.data["balance_calculado"])
